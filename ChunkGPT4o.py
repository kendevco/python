import os
import shutil
import fitz  # PyMuPDF
import base64
import openai
import logging
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Set up logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger()

# Set your OpenAI API key
openai.api_key = os.getenv('OPENAI_API_KEY') or 'YOUR_OPENAI_API_KEY'

# Directory for temporary files
TEMP_DIR = "temp"
OUTPUT_DOCX = os.path.join(TEMP_DIR, "output.docx")
CACHE_DIR = os.path.join(TEMP_DIR, "cache")

# Ensure the temp and cache directories exist
os.makedirs(TEMP_DIR, exist_ok=True)
os.makedirs(CACHE_DIR, exist_ok=True)

# Function to encode image to base64
def encode_image(image_path):
    try:
        with open(image_path, "rb") as image_file:
            return base64.b64encode(image_file.read()).decode('utf-8')
    except Exception as e:
        logger.error(f"Error encoding image {image_path}: {e}")
        return None

# Function to process PDF and save pages as images
def process_pdf(file_path):
    try:
        pdf_document = fitz.open(file_path)
        output_images = []

        for page_num in range(len(pdf_document)):
            img_path = os.path.join(TEMP_DIR, f"{os.path.basename(file_path)}_page_{page_num + 1}.png")
            if os.path.exists(img_path):
                logger.info(f"Skipping existing file: {img_path}")
                output_images.append(img_path)
                continue
            page = pdf_document.load_page(page_num)
            pix = page.get_pixmap()
            pix.save(img_path)
            output_images.append(img_path)
            logger.info(f"Saved image: {img_path}")

        return output_images
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {e}")
        return []

# Function to create messages for images
def create_image_messages(image_files):
    images_content = [{"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encode_image(img)}"}} for img in image_files if encode_image(img) is not None]
    if not images_content:
        logger.warning("No valid images to process")
        return []

    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": "Please extract all the text from these images. Correct minor mistakes and interpret the author's intent where letters may look ambiguous. Provide detailed and complete text."}
            ] + images_content
        }
    ]
    return messages

# Function to split list into batches
def split_into_batches(lst, batch_size):
    for i in range(0, len(lst), batch_size):
        yield lst[i:i + batch_size]

# Function to process and analyze images in batches
def analyze_images(image_files, batch_size=5, max_batches=None):
    results = []
    batches = list(split_into_batches(image_files, batch_size))
    
    # Limit the number of batches if max_batches is specified
    if max_batches is not None:
        batches = batches[:max_batches]
    
    for i, batch in enumerate(batches):
        messages = create_image_messages(batch)
        if not messages:
            continue

        cache_miss = False
        for img_path in batch:
            cache_path = os.path.join(CACHE_DIR, os.path.basename(img_path) + ".txt")
            if not os.path.exists(cache_path):
                cache_miss = True
                break

        if cache_miss:
            logger.info(f"Processing batch {i + 1}/{len(batches)} with OpenAI")
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-4o",
                    messages=messages,
                    max_tokens=2500  # Increased token limit for more detailed responses
                )
                logger.info(f"Received response for batch {i + 1}")
                batch_result = response.choices[0]['message']['content']

                for img_path in batch:
                    cache_path = os.path.join(CACHE_DIR, os.path.basename(img_path) + ".txt")
                    with open(cache_path, "w") as cache_file:
                        cache_file.write(batch_result)
            except Exception as e:
                logger.error(f"Error processing images: {e}")
                batch_result = "Error processing images."
        else:
            logger.info(f"Cache hit for batch {i + 1}")
            batch_result = ""
            for img_path in batch:
                cache_path = os.path.join(CACHE_DIR, os.path.basename(img_path) + ".txt")
                with open(cache_path, "r") as cache_file:
                    batch_result += cache_file.read() + "\n"

        results.append((batch, batch_result))
    
    return results

# Function to process directory and generate a .docx file
def process_directory(directory, output_docx, batch_size=5, max_batches=None, cleanup=True):
    document = Document()

    # Ensure file overwriting
    if os.path.exists(output_docx):
        os.remove(output_docx)

    for filename in os.listdir(directory):
        file_path = os.path.join(directory, filename)
        try:
            if filename.endswith('.pdf'):
                logger.info(f"Processing PDF: {filename}")
                image_files = process_pdf(file_path)
                analysis_results = analyze_images(image_files, batch_size=batch_size, max_batches=max_batches)

                for batch, analysis_result in analysis_results:
                    analysis_result_lines = analysis_result.split('\n')
                    current_page = 0
                    page_texts = {}

                    for line in analysis_result_lines:
                        if line.strip().startswith("**Image "):
                            current_page += 1
                            page_texts[current_page] = []
                        elif current_page > 0:
                            page_texts[current_page].append(line)

                    if not page_texts:
                        logger.warning(f"No text found for batch: {batch}")

                    for img_path in batch:
                        try:
                            page_number = image_files.index(img_path) + 1
                            heading = f"{os.path.basename(file_path)} - Page {page_number}"
                            document.add_heading(heading, level=2)
                            
                            paragraph = document.add_paragraph()
                            run = paragraph.add_run()
                            run.add_picture(img_path, width=Inches(3.75))  # Centered image at 50% scale
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                            text_content = "\n".join(page_texts.get(page_number, ["No text found for this image."]))
                            paragraph = document.add_paragraph(text_content)
                            paragraph_format = paragraph.paragraph_format
                            paragraph_format.left_indent = Inches(0.5)
                            paragraph_format.right_indent = Inches(0.5)
                            paragraph_format.space_before = Inches(0.5)
                            paragraph_format.space_after = Inches(0.5)

                            # Adding a page break after each page's content
                            document.add_page_break()
                        except Exception as e:
                            logger.error(f"Error adding content for image {img_path}: {e}")
            else:
                logger.info(f"Skipping non-PDF file: {filename}")
        except Exception as e:
            logger.error(f"Error processing {filename}: {e}")

    try:
        document.save(output_docx)
        logger.info(f"Document saved as {output_docx}")
    except Exception as e:
        logger.error(f"Error saving document {output_docx}: {e}")

    # Clean up temporary files if specified
    if cleanup:
        try:
            shutil.rmtree(TEMP_DIR)
            logger.info(f"Temporary files cleaned up")
        except Exception as e:
            logger.error(f"Error cleaning up temporary files: {e}")

if __name__ == "__main__":
    # Test with specific batch size and number of batches, save to temp\output.docx, and do not cleanup temp images
    process_directory(r"C:\Users\kenne\Videos\Captures\Process", OUTPUT_DOCX, batch_size=5, max_batches=2, cleanup=False)
